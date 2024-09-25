import unittest
from typing import Optional, Dict

import docx_ea_font
from docx import Document
from docx.oxml.ns import qn


def find_with_optional_ns(element, path: str, namespaces: Optional[Dict[str, str]] = None):
    return element.find(path, namespaces=namespaces)


class TestSetFont(unittest.TestCase):
    def test_set_font(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("测试文本")

        docx_ea_font.set_font(run, "微软雅黑")

        self.assertEqual(run.font.name, "微软雅黑")
        self.assertIsNotNone(run._element.rPr)
        self.assertIsNotNone(run._element.rPr.rFonts)
        self.assertEqual(run._element.rPr.rFonts.get(qn("w:eastAsia")), "微软雅黑")

    def test_change_font_to_kaiti(self):
        doc = Document("tests/test.docx")
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                docx_ea_font.set_font(run, "楷体")

        doc.save("test_kaiti.docx")


if __name__ == "__main__":
    unittest.main()
